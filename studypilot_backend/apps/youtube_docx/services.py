import logging
import re
import uuid
from datetime import timedelta
from html import unescape
from pathlib import Path
from urllib.parse import parse_qs, urlparse

import requests
from django.conf import settings
from django.core.exceptions import ImproperlyConfigured
from django.utils import timezone

from apps.ai.services import AIServiceError, generate_structured_docx_content_with_deepseek
from apps.documents.services import clean_extracted_text, clean_safe_string

logger = logging.getLogger(__name__)


class YouTubeDocxError(Exception):
    pass


TRANSCRIPT_UNAVAILABLE_MESSAGE = "Transcript is not available for this video."


def canonical_youtube_url(video_id):
    return f"https://www.youtube.com/watch?v={video_id}"


def extract_youtube_video_id(youtube_url):
    url = (youtube_url or "").strip()
    if not re.match(r"^https?://", url, re.IGNORECASE):
        url = f"https://{url}"
    parsed = urlparse(url)
    host = parsed.netloc.lower().replace("www.", "")
    video_id = ""

    if host in {"youtube.com", "m.youtube.com", "music.youtube.com"}:
        if parsed.path == "/watch":
            video_id = (parse_qs(parsed.query).get("v") or [""])[0]
        for prefix in ("/shorts/", "/embed/"):
            if parsed.path.startswith(prefix):
                video_id = parsed.path.split("/")[2]
    if host == "youtu.be":
        video_id = parsed.path.strip("/").split("/")[0]

    video_id = re.sub(r"[^A-Za-z0-9_-]", "", video_id or "")
    if not re.fullmatch(r"[A-Za-z0-9_-]{11}", video_id):
        return ""
    return video_id


parse_youtube_video_id = extract_youtube_video_id


def format_duration(seconds):
    total = int(seconds or 0)
    hours, remainder = divmod(total, 3600)
    minutes, secs = divmod(remainder, 60)
    if hours:
        return f"{hours}:{minutes:02d}:{secs:02d}"
    return f"{minutes}:{secs:02d}"


def safe_filename(value, fallback="youtube_lecture"):
    cleaned = clean_safe_string(value, fallback=fallback, max_length=90).lower()
    cleaned = re.sub(r"[^a-z0-9]+", "_", cleaned).strip("_")
    return cleaned or fallback


def ensure_temp_dir():
    temp_dir = Path(settings.YOUTUBE_DOCX_TEMP_DIR)
    temp_dir.mkdir(parents=True, exist_ok=True)
    return temp_dir


def ensure_audio_temp_dir():
    temp_dir = Path(settings.YOUTUBE_AUDIO_TEMP_DIR)
    temp_dir.mkdir(parents=True, exist_ok=True)
    return temp_dir


def _delete_temp_path(path):
    if not path:
        return
    try:
        Path(path).unlink(missing_ok=True)
    except Exception:
        logger.warning("YouTube DOCX: could not delete temporary file")


def cleanup_old_docx_files():
    temp_dir = ensure_temp_dir()
    expiry = timezone.now() - timedelta(minutes=settings.YOUTUBE_DOCX_EXPIRY_MINUTES)
    for path in temp_dir.glob("*.docx"):
        try:
            modified = timezone.datetime.fromtimestamp(path.stat().st_mtime, tz=timezone.get_current_timezone())
            if modified < expiry:
                path.unlink(missing_ok=True)
        except Exception:
            logger.warning("Could not clean old YouTube DOCX temp file: %s", path, exc_info=True)

    audio_dir = ensure_audio_temp_dir()
    for path in audio_dir.glob("*"):
        if not path.is_file():
            continue
        try:
            modified = timezone.datetime.fromtimestamp(path.stat().st_mtime, tz=timezone.get_current_timezone())
            if modified < expiry:
                path.unlink(missing_ok=True)
        except Exception:
            logger.warning("Could not clean old YouTube audio temp file.", exc_info=True)


def get_temp_docx_path(temp_file_id):
    temp_dir = ensure_temp_dir()
    safe_id = re.sub(r"[^a-f0-9]", "", temp_file_id or "")
    if not safe_id:
        return None
    matches = list(temp_dir.glob(f"{safe_id}__*.docx"))
    return matches[0] if matches else None


def _transcript_items_to_text(items):
    lines = []
    for item in items:
        text = item.get("text") if isinstance(item, dict) else getattr(item, "text", "")
        if text:
            lines.append(text)
    cleaned = clean_extracted_text(" ".join(lines))
    return cleaned


def _validate_transcript_text(text):
    cleaned = clean_extracted_text(text)
    if len(cleaned) < 120:
        raise YouTubeDocxError(TRANSCRIPT_UNAVAILABLE_MESSAGE)
    return cleaned


def _fetch_transcript_object(transcript, label):
    try:
        text = _validate_transcript_text(_transcript_items_to_text(transcript.fetch()))
        logger.info("YouTube DOCX: youtube-transcript-api transcript found method=%s", label)
        return text
    except Exception:
        logger.info("YouTube DOCX: youtube-transcript-api transcript attempt failed method=%s", label)
        return ""


def _fetch_transcript_api_text(video_id):
    from youtube_transcript_api import YouTubeTranscriptApi

    logger.info("YouTube DOCX: youtube-transcript-api tried video_id=%s", video_id)
    try:
        api = YouTubeTranscriptApi()
        transcript_list = api.list(video_id)
    except AttributeError:
        transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
    try:
        transcripts = list(transcript_list)
    except Exception as exc:
        logger.info("YouTube DOCX: youtube-transcript-api list failed video_id=%s", video_id)
        raise YouTubeDocxError(TRANSCRIPT_UNAVAILABLE_MESSAGE) from exc

    for language in ("en", "en-US", "en-GB"):
        try:
            text = _fetch_transcript_object(transcript_list.find_manually_created_transcript([language]), f"manual_{language}")
            if text:
                return text
        except Exception:
            continue

    for language in ("en", "en-US", "en-GB"):
        try:
            text = _fetch_transcript_object(transcript_list.find_generated_transcript([language]), f"generated_{language}")
            if text:
                return text
        except Exception:
            continue

    for language in ("en", "en-US", "en-GB"):
        try:
            text = _fetch_transcript_object(transcript_list.find_transcript([language]), f"any_{language}")
            if text:
                return text
        except Exception:
            continue

    for transcript in transcripts:
        text = _fetch_transcript_object(transcript, "first_available")
        if text:
            return text

    for transcript in transcripts:
        try:
            if transcript.is_translatable:
                text = _fetch_transcript_object(transcript.translate("en"), "translated_to_en")
                if text:
                    return text
        except Exception:
            continue

    raise YouTubeDocxError(TRANSCRIPT_UNAVAILABLE_MESSAGE)


def _dedupe_caption_lines(text):
    lines = []
    previous = ""
    for line in re.split(r"[\r\n]+", text or ""):
        cleaned = clean_extracted_text(line)
        if not cleaned:
            continue
        normalized = re.sub(r"\W+", "", cleaned.lower())
        if normalized and normalized != previous:
            lines.append(cleaned)
            previous = normalized
    return " ".join(lines)


def _clean_json3_subtitle(raw_text):
    try:
        import json

        payload = json.loads(raw_text or "{}")
        lines = []
        for event in payload.get("events", []) or []:
            parts = []
            for segment in event.get("segs", []) or []:
                value = segment.get("utf8", "")
                if value:
                    parts.append(value)
            if parts:
                lines.append("".join(parts))
        return clean_extracted_text(_dedupe_caption_lines("\n".join(lines)))
    except Exception:
        return ""


def _clean_subtitle_text(raw_text):
    json3_text = _clean_json3_subtitle(raw_text)
    if json3_text:
        return json3_text

    try:
        import io
        import webvtt

        captions = webvtt.read_buffer(io.StringIO(raw_text or ""))
        parsed = _dedupe_caption_lines("\n".join(caption.text for caption in captions if caption.text))
        if parsed:
            return clean_extracted_text(parsed)
    except Exception:
        pass

    text = raw_text or ""
    text = re.sub(r"WEBVTT.*?(?=\n)", "", text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r"Kind:.*?(?=\n)", "", text, flags=re.IGNORECASE)
    text = re.sub(r"Language:.*?(?=\n)", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\d{1,2}:\d{2}:\d{2}[,.]\d{3}\s+-->\s+\d{1,2}:\d{2}:\d{2}[,.]\d{3}.*", "", text)
    text = re.sub(r"\d{2}:\d{2}[,.]\d{3}\s+-->\s+\d{2}:\d{2}[,.]\d{3}.*", "", text)
    text = re.sub(r"^\d+$", "", text, flags=re.MULTILINE)
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"&nbsp;|&amp;|&quot;|&#39;", " ", text)
    text = re.sub(r"\[[^\]]+\]", " ", text)
    text = _dedupe_caption_lines(text)
    text = re.sub(r"\s+", " ", text)
    return clean_extracted_text(text)


def _subtitle_candidates(info, include_automatic=False):
    requested = info.get("requested_subtitles") or {}
    subtitles = info.get("subtitles") or {}
    automatic = info.get("automatic_captions") or {}
    preferred = ["en", "en-US", "en-GB"]
    ordered = []
    seen = set()
    sources = [requested, subtitles]
    if include_automatic:
        sources.append(automatic)
    for source in sources:
        for language in preferred:
            for entry in source.get(language, []) or []:
                key = entry.get("url") or repr(entry)
                if key not in seen:
                    ordered.append(entry)
                    seen.add(key)
        for language, entries in source.items():
            for entry in entries or []:
                key = entry.get("url") or f"{language}:{repr(entry)}"
                if key not in seen:
                    ordered.append(entry)
                    seen.add(key)
    return ordered


def _fetch_ytdlp_subtitle_text(youtube_url, include_automatic=False):
    from yt_dlp import YoutubeDL

    source_label = "automatic captions" if include_automatic else "subtitles"
    logger.info("YouTube DOCX: trying yt-dlp %s", source_label)
    options = {
        "quiet": True,
        "no_warnings": True,
        "skip_download": True,
        "noplaylist": True,
        "writesubtitles": True,
        "writeautomaticsub": True,
        "subtitleslangs": ["en", "en-US", "en-GB", "all"],
    }
    with YoutubeDL(options) as ydl:
        info = ydl.extract_info(youtube_url, download=False)

    for entry in _subtitle_candidates(info, include_automatic=include_automatic):
        subtitle_url = entry.get("url")
        if not subtitle_url:
            continue
        try:
            response = requests.get(subtitle_url, timeout=12)
            response.raise_for_status()
            text = _validate_transcript_text(_clean_subtitle_text(response.text))
            logger.info("YouTube DOCX: transcript found yes source=yt_dlp_%s", "automatic_captions" if include_automatic else "subtitles")
            return text
        except Exception as exc:
            logger.info("YouTube DOCX: yt-dlp %s candidate failed reason=%s", source_label, exc.__class__.__name__)
            continue
    logger.info("YouTube DOCX: yt-dlp %s failed reason=no_readable_caption_text", source_label)
    raise YouTubeDocxError(TRANSCRIPT_UNAVAILABLE_MESSAGE)


def _download_youtube_audio(youtube_url, video_id):
    from yt_dlp import YoutubeDL

    logger.info("YouTube DOCX: trying audio transcription download video_id=%s", video_id)
    audio_dir = ensure_audio_temp_dir()
    output_template = str(audio_dir / f"{video_id}_{uuid.uuid4().hex}.%(ext)s")
    options = {
        "quiet": True,
        "no_warnings": True,
        "format": "bestaudio[ext=m4a]/bestaudio[ext=webm]/bestaudio",
        "outtmpl": output_template,
        "noplaylist": True,
        "socket_timeout": 20,
        "retries": 2,
        "fragment_retries": 2,
    }
    before = {path.resolve() for path in audio_dir.glob("*") if path.is_file()}
    with YoutubeDL(options) as ydl:
        info = ydl.extract_info(youtube_url, download=True)
        downloaded = ydl.prepare_filename(info)
    candidate = Path(downloaded)
    if candidate.exists():
        return candidate
    after = [path for path in audio_dir.glob("*") if path.is_file() and path.resolve() not in before]
    if after:
        return max(after, key=lambda path: path.stat().st_mtime)
    raise YouTubeDocxError(TRANSCRIPT_UNAVAILABLE_MESSAGE)


def transcribe_youtube_audio(youtube_url):
    video_id = extract_youtube_video_id(youtube_url)
    if not video_id:
        raise YouTubeDocxError("Invalid YouTube link.")
    audio_path = None
    try:
        audio_path = _download_youtube_audio(canonical_youtube_url(video_id), video_id)
        return _transcribe_audio_with_faster_whisper(audio_path)
    finally:
        if audio_path:
            _delete_temp_path(audio_path)


def _transcribe_audio_with_faster_whisper(audio_path):
    try:
        from faster_whisper import WhisperModel
    except Exception as exc:
        logger.info("YouTube DOCX: audio transcription unavailable reason=missing_faster_whisper")
        raise YouTubeDocxError(TRANSCRIPT_UNAVAILABLE_MESSAGE) from exc

    try:
        logger.info("YouTube DOCX: Whisper transcription started")
        model = WhisperModel(settings.WHISPER_MODEL_SIZE, device="cpu", compute_type="int8")
        segments, _ = model.transcribe(str(audio_path), beam_size=3, vad_filter=True)
        text = " ".join(segment.text.strip() for segment in segments if getattr(segment, "text", "").strip())
        transcript = _validate_transcript_text(text)
        logger.info("YouTube DOCX: Whisper transcription completed transcript_length=%s", len(transcript))
        return transcript
    except Exception as exc:
        logger.info("YouTube DOCX: audio transcription failed reason=%s", exc.__class__.__name__)
        raise YouTubeDocxError(TRANSCRIPT_UNAVAILABLE_MESSAGE) from exc


def _fetch_audio_transcript_text(youtube_url, video_id):
    if not settings.ENABLE_AUDIO_TRANSCRIPTION:
        logger.info("YouTube DOCX: ENABLE_AUDIO_TRANSCRIPTION is false")
        raise YouTubeDocxError(TRANSCRIPT_UNAVAILABLE_MESSAGE)

    logger.info("YouTube DOCX: ENABLE_AUDIO_TRANSCRIPTION is true")
    logger.info("YouTube DOCX: trying audio transcription video_id=%s", video_id)
    audio_path = None
    try:
        audio_path = _download_youtube_audio(youtube_url, video_id)
        logger.info("YouTube DOCX: temporary audio downloaded yes video_id=%s", video_id)
        transcript = _transcribe_audio_with_faster_whisper(audio_path)
        logger.info("YouTube DOCX: transcript found yes source=audio_transcription")
        return transcript
    finally:
        if audio_path:
            _delete_temp_path(audio_path)


def fetch_youtube_transcript(video_id):
    try:
        return _fetch_transcript_api_text(video_id)
    except YouTubeDocxError:
        raise
    except Exception as exc:
        raise YouTubeDocxError(TRANSCRIPT_UNAVAILABLE_MESSAGE) from exc


def fetch_youtube_transcript_with_fallback(video_id, youtube_url, allow_audio=True):
    try:
        text = fetch_youtube_transcript(video_id)
        logger.info("YouTube DOCX: transcript found yes source=youtube_transcript_api")
        return text
    except YouTubeDocxError as exc:
        logger.info("YouTube DOCX: youtube-transcript-api failed reason=%s", exc.__class__.__name__)
    try:
        return _fetch_ytdlp_subtitle_text(youtube_url, include_automatic=False)
    except YouTubeDocxError as exc:
        logger.info("YouTube DOCX: yt-dlp subtitles failed reason=%s", exc.__class__.__name__)
    except Exception as exc:
        logger.info("YouTube DOCX: yt-dlp subtitles failed reason=%s", exc.__class__.__name__)
    try:
        return _fetch_ytdlp_subtitle_text(youtube_url, include_automatic=True)
    except YouTubeDocxError as exc:
        logger.info("YouTube DOCX: yt-dlp automatic captions failed reason=%s", exc.__class__.__name__)
    except Exception as exc:
        logger.info("YouTube DOCX: yt-dlp automatic captions failed reason=%s", exc.__class__.__name__)
    if not allow_audio:
        raise YouTubeDocxError(TRANSCRIPT_UNAVAILABLE_MESSAGE)
    try:
        return _fetch_audio_transcript_text(youtube_url, video_id)
    except YouTubeDocxError as exc:
        if settings.ENABLE_AUDIO_TRANSCRIPTION:
            logger.info("YouTube DOCX: audio transcription failed; manual transcript required.")
            raise YouTubeDocxError("Audio transcription failed. Paste transcript manually to continue.") from exc
        raise
    except Exception as exc:
        raise YouTubeDocxError(TRANSCRIPT_UNAVAILABLE_MESSAGE) from exc


def get_youtube_transcript(youtube_url):
    video_id = extract_youtube_video_id(youtube_url)
    if not video_id:
        raise YouTubeDocxError("Invalid YouTube link.")
    return fetch_youtube_transcript_with_fallback(video_id, canonical_youtube_url(video_id), allow_audio=True)


def fetch_youtube_metadata(youtube_url, video_id):
    logger.info("YouTube DOCX: metadata lookup started video_id=%s", video_id)
    metadata = {
        "title": f"YouTube Lecture {video_id}",
        "channel": "",
        "duration": 0,
        "thumbnail": "",
        "url": canonical_youtube_url(video_id),
        "webpage_url": canonical_youtube_url(video_id),
        "video_id": video_id,
    }
    try:
        from yt_dlp import YoutubeDL

        with YoutubeDL({"quiet": True, "no_warnings": True, "skip_download": True, "noplaylist": True, "socket_timeout": 12}) as ydl:
            info = ydl.extract_info(youtube_url, download=False)
        metadata.update(
            {
                "title": clean_safe_string(info.get("title"), fallback=metadata["title"], max_length=180),
                "channel": clean_safe_string(info.get("channel") or info.get("uploader"), max_length=140),
                "duration": int(info.get("duration") or 0),
                "thumbnail": info.get("thumbnail") or "",
                "url": canonical_youtube_url(video_id),
                "webpage_url": info.get("webpage_url") or canonical_youtube_url(video_id),
                "video_id": video_id,
            }
        )
        logger.info("YouTube DOCX: metadata fetched yes video_id=%s source=yt_dlp", video_id)
        return metadata
    except Exception:
        logger.warning("YouTube DOCX: yt-dlp metadata lookup failed video_id=%s; trying public metadata fallback.", video_id, exc_info=True)

    try:
        response = requests.get(
            "https://www.youtube.com/oembed",
            params={"url": canonical_youtube_url(video_id), "format": "json"},
            timeout=10,
        )
        response.raise_for_status()
        info = response.json()
        metadata.update(
            {
                "title": clean_safe_string(info.get("title"), fallback=metadata["title"], max_length=180),
                "channel": clean_safe_string(info.get("author_name"), max_length=140),
                "thumbnail": info.get("thumbnail_url") or f"https://img.youtube.com/vi/{video_id}/hqdefault.jpg",
                "url": canonical_youtube_url(video_id),
                "webpage_url": canonical_youtube_url(video_id),
                "video_id": video_id,
            }
        )
        logger.info("YouTube DOCX: metadata fetched yes video_id=%s source=oembed", video_id)
    except Exception:
        metadata["thumbnail"] = f"https://img.youtube.com/vi/{video_id}/hqdefault.jpg"
        logger.warning("YouTube DOCX: metadata fetched no video_id=%s; using safe fallback.", video_id, exc_info=True)
    return metadata


def analyze_youtube_video(youtube_url):
    cleanup_old_docx_files()
    logger.info("YouTube DOCX: analyze request received")
    logger.info("YouTube DOCX: URL received %s", "yes" if youtube_url else "no")
    video_id = extract_youtube_video_id(youtube_url)
    if not video_id:
        logger.info("YouTube DOCX: video ID extracted no")
        raise YouTubeDocxError("Invalid YouTube link.")
    logger.info("YouTube DOCX: video ID extracted yes video_id=%s", video_id)

    canonical_url = canonical_youtube_url(video_id)
    metadata = fetch_youtube_metadata(canonical_url, video_id)
    has_transcript = True
    try:
        fetch_youtube_transcript_with_fallback(video_id, canonical_url, allow_audio=False)
        logger.info("YouTube DOCX: transcript fetched yes video_id=%s", video_id)
    except YouTubeDocxError:
        has_transcript = False
        logger.info("YouTube DOCX: transcript fetched no video_id=%s", video_id)

    return {
        "video_id": video_id,
        "title": metadata.get("title") or f"YouTube Lecture {video_id}",
        "channel": metadata.get("channel") or "Unknown channel",
        "duration": format_duration(metadata.get("duration")),
        "duration_seconds": int(metadata.get("duration") or 0),
        "thumbnail": metadata.get("thumbnail") or "",
        "youtube_url": canonical_url,
        "has_transcript": has_transcript,
        "manual_transcript_required": not has_transcript,
    }


def estimate_pages(transcript, metadata, detail_level="comprehensive"):
    duration = int(metadata.get("duration") or 0)
    multipliers = {
        "summary": 0.7,
        "comprehensive": 1.0,
        "full_study_notes": 1.25,
    }
    multiplier = multipliers.get(detail_level, 1.0)
    by_duration = int((duration // 240) * multiplier) if duration else 0
    by_text = int(max(0, len(transcript or "") // 1800) * multiplier)
    return max(30, min(55, max(by_duration, by_text, 30)))


def _set_document_styles(document):
    from docx.shared import Pt, RGBColor

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    for style_name, size, color in (
        ("Heading 1", 18, RGBColor(15, 23, 42)),
        ("Heading 2", 15, RGBColor(37, 99, 235)),
        ("Heading 3", 13, RGBColor(15, 23, 42)),
    ):
        if style_name in styles:
            styles[style_name].font.name = "Aptos"
            styles[style_name].font.size = Pt(size)
            styles[style_name].font.bold = True
            styles[style_name].font.color.rgb = color
    for section in document.sections:
        section.header.paragraphs[0].text = "StudyPilot YouTube Study Document"
        section.footer.paragraphs[0].text = "Generated by StudyPilot"


def clean_docx_text(text, keep_pipes=False, preserve_emphasis=False):
    """Clean model/local text before it is written to a Word paragraph."""
    if text is None:
        return ""
    cleaned = unescape(str(text))
    cleaned = cleaned.replace("\x00", "").replace("\u0000", "").replace("\ufeff", "")
    cleaned = re.sub(r"\\([\\`*_{}\[\]()#+\-.!|])", r"\1", cleaned)
    cleaned = cleaned.replace("\\", "")
    cleaned = re.sub(r"\$\$?", "", cleaned)
    cleaned = re.sub(r"(?<!:)//+", " ", cleaned)
    cleaned = re.sub(r"`+", "", cleaned)
    cleaned = re.sub(r"^\s*#{1,6}\s*", "", cleaned)
    cleaned = re.sub(r"^\s*>+\s*", "", cleaned)
    cleaned = re.sub(r"^\s*[-*_]{3,}\s*$", "", cleaned)
    cleaned = re.sub(r"<[^>]+>", "", cleaned)
    if not preserve_emphasis:
        cleaned = re.sub(r"\*\*(.*?)\*\*", r"\1", cleaned)
        cleaned = re.sub(r"__(.*?)__", r"\1", cleaned)
        cleaned = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", cleaned)
        cleaned = re.sub(r"(?<!_)_([^_]+)_(?!_)", r"\1", cleaned)
    if not keep_pipes:
        cleaned = cleaned.replace("|", " ")
    cleaned = clean_extracted_text(cleaned)
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    return cleaned.strip()


def _add_heading(document, text, level=1):
    paragraph = document.add_heading(clean_docx_text(text), level=level)
    for run in paragraph.runs:
        run.font.name = "Calibri"
    return paragraph


def _add_formatted_paragraph(document, text, style=None):
    from docx.shared import Pt

    cleaned = clean_docx_text(text, preserve_emphasis=True)
    if not cleaned:
        return None

    paragraph = document.add_paragraph(style=style) if style else document.add_paragraph()
    pattern = re.compile(r"(\*\*.+?\*\*|__.+?__|\*[^*]+\*|_[^_]+_)")
    position = 0
    for match in pattern.finditer(cleaned):
        before = clean_docx_text(cleaned[position : match.start()])
        if before:
            paragraph.add_run(before)
        token = match.group(0)
        is_bold = token.startswith("**") or token.startswith("__")
        stripped = token[2:-2] if is_bold else token[1:-1]
        text_run = clean_docx_text(stripped)
        if text_run:
            run = paragraph.add_run(text_run)
            run.bold = is_bold
            run.italic = not is_bold
        position = match.end()
    after = clean_docx_text(cleaned[position:])
    if after:
        paragraph.add_run(after)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def _is_table_separator(cells):
    return bool(cells) and all(re.fullmatch(r"[:\-\s]+", cell or "") for cell in cells)


def _split_table_row(line):
    cells = [clean_docx_text(cell, keep_pipes=False) for cell in line.strip().strip("|").split("|")]
    return [cell for cell in cells if cell]


def _add_table_or_fallback(document, table_lines):
    from docx.shared import Pt

    rows = []
    for line in table_lines:
        cells = _split_table_row(line)
        if not cells or _is_table_separator(cells):
            continue
        rows.append(cells)

    max_cols = max((len(row) for row in rows), default=0)
    if len(rows) < 2 or max_cols < 2:
        for line in table_lines:
            clean_line = clean_docx_text(line)
            if clean_line:
                document.add_paragraph(clean_line, style="List Bullet")
        return

    normalized_rows = [row + [""] * (max_cols - len(row)) for row in rows]
    table = document.add_table(rows=1, cols=max_cols)
    table.style = "Table Grid"
    for index, cell_text in enumerate(normalized_rows[0]):
        cell = table.rows[0].cells[index]
        cell.text = ""
        run = cell.paragraphs[0].add_run(cell_text)
        run.bold = True
    for row in normalized_rows[1:]:
        cells = table.add_row().cells
        for index, cell_text in enumerate(row):
            cells[index].text = cell_text
    document.add_paragraph().paragraph_format.space_after = Pt(4)


def _add_markdown_content(document, content):
    table_buffer = []
    major_heading_seen = False

    def flush_table():
        nonlocal table_buffer
        if table_buffer:
            _add_table_or_fallback(document, table_buffer)
            table_buffer = []

    for raw_line in (content or "").splitlines():
        line = raw_line.strip()
        if not line:
            flush_table()
            continue
        if re.fullmatch(r"[-*_]{3,}", line):
            flush_table()
            continue
        if line.count("|") >= 2:
            table_buffer.append(line)
            continue

        flush_table()
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            if level == 1 and major_heading_seen:
                document.add_page_break()
            major_heading_seen = major_heading_seen or level == 1
            _add_heading(document, heading_match.group(2), level)
        elif line.startswith(("- ", "* ")):
            _add_formatted_paragraph(document, line[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            _add_formatted_paragraph(document, re.sub(r"^\d+\.\s+", "", line), style="List Number")
        else:
            _add_formatted_paragraph(document, line)
    flush_table()


def _as_clean_list(value):
    if not value:
        return []
    if isinstance(value, list):
        items = value
    else:
        items = re.split(r"[\r\n]+", str(value))
    cleaned = []
    for item in items:
        if isinstance(item, dict):
            text = item.get("text") or item.get("title") or item.get("value") or ""
        else:
            text = item
        text = clean_docx_text(text)
        if text:
            cleaned.append(text)
    return cleaned


def _add_bullet_list(document, items):
    for item in _as_clean_list(items):
        _add_formatted_paragraph(document, item, style="List Bullet")


def _add_numbered_list(document, items):
    for item in _as_clean_list(items):
        _add_formatted_paragraph(document, item, style="List Number")


def _add_two_column_table(document, headers, rows):
    cleaned_rows = []
    for left, right in rows:
        left_text = clean_docx_text(left)
        right_text = clean_docx_text(right)
        if left_text or right_text:
            cleaned_rows.append((left_text, right_text))
    if not cleaned_rows:
        return
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = ""
        run = cell.paragraphs[0].add_run(clean_docx_text(header))
        run.bold = True
    for left, right in cleaned_rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
    document.add_paragraph()


def _add_structured_content(document, content):
    if not isinstance(content, dict):
        _add_markdown_content(document, content)
        return

    introduction = clean_docx_text(content.get("introduction", ""))
    if introduction:
        _add_heading(document, "Introduction", 1)
        _add_formatted_paragraph(document, introduction)
        document.add_page_break()

    objectives = _as_clean_list(content.get("learning_objectives"))
    if objectives:
        _add_heading(document, "Learning Objectives", 1)
        _add_bullet_list(document, objectives)
        document.add_page_break()

    sections = content.get("sections") if isinstance(content.get("sections"), list) else []
    if sections:
        _add_heading(document, "Main Study Notes", 1)
        for index, section in enumerate(sections, 1):
            if not isinstance(section, dict):
                continue
            heading = clean_docx_text(section.get("heading") or f"Lecture Section {index}")
            _add_heading(document, heading, 2)
            summary = clean_docx_text(section.get("summary", ""))
            if summary:
                _add_formatted_paragraph(document, summary)
            key_points = _as_clean_list(section.get("key_points"))
            if key_points:
                _add_heading(document, "Key Points", 3)
                _add_bullet_list(document, key_points)
            examples = _as_clean_list(section.get("examples"))
            if examples:
                _add_heading(document, "Examples", 3)
                _add_bullet_list(document, examples)
        document.add_page_break()

    key_concepts = content.get("key_concepts") if isinstance(content.get("key_concepts"), list) else []
    if key_concepts:
        _add_heading(document, "Key Concepts and Definitions", 1)
        rows = []
        for item in key_concepts:
            if isinstance(item, dict):
                rows.append((item.get("term", ""), item.get("definition", "")))
        _add_two_column_table(document, ("Term", "Definition"), rows)
        document.add_page_break()

    takeaways = _as_clean_list(content.get("important_takeaways"))
    if takeaways:
        _add_heading(document, "Important Takeaways", 1)
        _add_bullet_list(document, takeaways)
        document.add_page_break()

    summary = clean_docx_text(content.get("summary", ""))
    if summary:
        _add_heading(document, "Summary", 1)
        _add_formatted_paragraph(document, summary)
        document.add_page_break()

    revision_questions = content.get("revision_questions") if isinstance(content.get("revision_questions"), list) else []
    if revision_questions:
        _add_heading(document, "Revision Questions", 1)
        for index, item in enumerate(revision_questions, 1):
            if isinstance(item, dict):
                question = clean_docx_text(item.get("question", ""))
                answer = clean_docx_text(item.get("answer", ""))
            else:
                question = clean_docx_text(item)
                answer = ""
            if question:
                _add_formatted_paragraph(document, f"{index}. {question}")
            if answer:
                _add_formatted_paragraph(document, f"Answer: {answer}")
        document.add_page_break()

    mcqs = content.get("mcqs") if isinstance(content.get("mcqs"), list) else []
    if mcqs:
        _add_heading(document, "MCQs with Answers", 1)
        for index, item in enumerate(mcqs, 1):
            if not isinstance(item, dict):
                continue
            question = clean_docx_text(item.get("question", ""))
            if question:
                _add_formatted_paragraph(document, f"{index}. {question}")
            _add_bullet_list(document, item.get("options"))
            correct = clean_docx_text(item.get("correct_answer", ""))
            explanation = clean_docx_text(item.get("explanation", ""))
            if correct:
                _add_formatted_paragraph(document, f"Correct answer: {correct}")
            if explanation:
                _add_formatted_paragraph(document, f"Explanation: {explanation}")
        document.add_page_break()

    glossary = content.get("glossary") if isinstance(content.get("glossary"), list) else []
    if glossary:
        _add_heading(document, "Glossary", 1)
        rows = []
        for item in glossary:
            if isinstance(item, dict):
                rows.append((item.get("term", ""), item.get("meaning", "")))
        _add_two_column_table(document, ("Term", "Meaning"), rows)
        document.add_page_break()

    checklist = _as_clean_list(content.get("study_checklist"))
    if checklist:
        _add_heading(document, "Final Study Checklist", 1)
        _add_bullet_list(document, checklist)


def _structured_sections_count(content):
    if not isinstance(content, dict):
        return len(re.findall(r"^#{1,3}\s+", content or "", flags=re.MULTILINE)) or 12
    count = 0
    for key in (
        "introduction",
        "learning_objectives",
        "sections",
        "key_concepts",
        "important_takeaways",
        "summary",
        "revision_questions",
        "mcqs",
        "glossary",
        "study_checklist",
    ):
        value = content.get(key)
        if value:
            count += len(value) if key == "sections" and isinstance(value, list) else 1
    return count or 12


def _add_transcript_reference(document, transcript):
    from docx.shared import Pt

    chunks = [chunk.strip() for chunk in re.split(r"(?<=[.!?])\s+", transcript) if chunk.strip()]
    if not chunks:
        return
    document.add_page_break()
    _add_heading(document, "Lecture Transcript Study Reference", 1)
    paragraph = document.add_paragraph(
        "These transcript-based reference notes preserve useful lecture wording so the generated study document stays grounded in the source video."
    )
    paragraph.paragraph_format.space_after = Pt(8)

    current = []
    section_index = 1
    for sentence in chunks[:420]:
        current.append(sentence)
        if len(" ".join(current)) >= 1200:
            _add_heading(document, f"Reference Segment {section_index}", 2)
            document.add_paragraph(" ".join(current))
            current = []
            section_index += 1
            if section_index > 18:
                break
    if current and section_index <= 18:
        _add_heading(document, f"Reference Segment {section_index}", 2)
        document.add_paragraph(" ".join(current))


def _document_label(value):
    labels = {
        "summary": "Summary",
        "comprehensive": "Comprehensive",
        "full_study_notes": "Full Study Notes",
        "study_guide": "Study Guide",
        "tutorial": "Tutorial",
        "lecture_notes": "Lecture Notes",
        "exam_revision": "Exam Revision",
    }
    return labels.get(value, value.replace("_", " ").title() if value else "")


def _add_video_snapshot_table(document, metadata, document_options):
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Video", metadata.get("title") or "YouTube Lecture"),
        ("Channel", metadata.get("channel") or "Unknown channel"),
        ("Source", metadata.get("url") or ""),
        ("Duration", format_duration(metadata.get("duration"))),
        ("Detail Level", _document_label(document_options.get("detail_level", "comprehensive"))),
        ("Document Style", _document_label(document_options.get("document_style", "study_guide"))),
        ("Key Frames / Markers", str(document_options.get("key_frames", 5))),
    ]
    custom = document_options.get("custom_instruction")
    if custom:
        rows.append(("Custom Focus", custom))
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)


def build_docx_file(metadata, transcript, generated_content, target_pages, document_options=None):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    document_options = document_options or {}
    temp_id = uuid.uuid4().hex
    content_title = generated_content.get("title") if isinstance(generated_content, dict) else ""
    title = clean_docx_text(content_title) or metadata.get("title") or "YouTube Lecture"
    filename = f"studypilot_youtube_notes_{safe_filename(title)}.docx"
    path = ensure_temp_dir() / f"{temp_id}__{filename}"

    document = Document()
    _set_document_styles(document)
    for section in document.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    _add_heading(document, title, 0)
    document.add_paragraph(f"Video source: {metadata.get('url') or ''}")
    document.add_paragraph(f"Channel: {metadata.get('channel') or 'Unknown channel'}")
    document.add_paragraph("Generated by StudyPilot")
    document.add_paragraph(f"Date generated: {timezone.now().date().isoformat()}")
    document.add_paragraph(
        "This document is a StudyPilot-generated academic study guide based on the available YouTube transcript. It is not a verbatim transcript."
    )
    if document_options.get("short_video_note"):
        note = document.add_paragraph(
            "This video transcript was short, so StudyPilot created the most complete document possible without adding filler."
        )
        note.runs[0].font.color.rgb = RGBColor(37, 99, 235)
    document.add_page_break()

    _add_heading(document, "Video and Document Options", 1)
    _add_video_snapshot_table(document, metadata, document_options)
    document.add_page_break()

    _add_heading(document, "Table of Contents", 1)
    for item in [
        "Introduction",
        "Learning Objectives",
        "Main Lecture Notes",
        "Key Concepts",
        "Definitions",
        "Examples and Step-by-Step Explanations",
        "Important Takeaways",
        "Summary",
        "Revision Questions",
        "MCQs with Answers",
        "Short Answer Questions",
        "Glossary",
        "Final Study Checklist",
    ]:
        document.add_paragraph(item, style="List Bullet")
    document.add_page_break()

    _add_structured_content(document, generated_content)

    for paragraph in document.paragraphs:
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.15

    document.save(path)
    return {
        "temp_file_id": temp_id,
        "filename": filename,
        "path": path,
        "estimated_pages": target_pages,
    }


def normalize_document_options(options):
    options = options or {}
    detail_level = (options.get("detail_level") or "comprehensive").strip().lower()
    document_style = (options.get("document_style") or "study_guide").strip().lower()
    if detail_level not in {"summary", "comprehensive", "full_study_notes"}:
        detail_level = "comprehensive"
    if document_style not in {"study_guide", "tutorial", "lecture_notes", "exam_revision"}:
        document_style = "study_guide"
    try:
        key_frames = int(options.get("key_frames", 5))
    except (TypeError, ValueError):
        key_frames = 5
    key_frames = max(0, min(12, key_frames))
    custom_instruction = clean_safe_string(options.get("custom_instruction", ""), max_length=300)
    return {
        "detail_level": detail_level,
        "document_style": document_style,
        "key_frames": key_frames,
        "custom_instruction": custom_instruction,
    }


def generate_youtube_docx(youtube_url, manual_transcript="", document_options=None):
    cleanup_old_docx_files()
    logger.info("YouTube DOCX generate called")
    logger.info("YouTube DOCX: URL received %s", "yes" if youtube_url else "no")
    video_id = extract_youtube_video_id(youtube_url)
    if not video_id:
        logger.info("YouTube DOCX: video ID extracted no")
        raise YouTubeDocxError("Invalid YouTube link.")
    logger.info("YouTube DOCX: video ID extracted yes video_id=%s", video_id)

    canonical_url = canonical_youtube_url(video_id)
    document_options = normalize_document_options(document_options)
    manual_text = clean_extracted_text(manual_transcript or "")
    if manual_text:
        transcript = _validate_transcript_text(manual_text)
        logger.info("YouTube DOCX: manual transcript used yes video_id=%s", video_id)
    else:
        logger.info("YouTube DOCX: manual transcript used no video_id=%s", video_id)
        try:
            transcript = fetch_youtube_transcript_with_fallback(video_id, canonical_url)
            logger.info("YouTube DOCX: transcript fetched yes video_id=%s", video_id)
        except YouTubeDocxError:
            logger.info("YouTube DOCX: transcript fetched no video_id=%s", video_id)
            raise
    metadata = fetch_youtube_metadata(canonical_url, video_id)
    target_pages = estimate_pages(transcript, metadata, detail_level=document_options["detail_level"])
    document_options["target_pages"] = target_pages
    sections_count = 12
    try:
        logger.info("YouTube DOCX: sending transcript to DeepSeek transcript_length=%s", len(transcript or ""))
        generated_content = generate_structured_docx_content_with_deepseek(transcript, metadata, document_options)
        sections_count = _structured_sections_count(generated_content)
        logger.info("YouTube DOCX: structured content generated yes video_id=%s", video_id)
    except ImproperlyConfigured:
        logger.info("YouTube DOCX: structured content generated no video_id=%s reason=api_key_unavailable", video_id)
        raise YouTubeDocxError("DeepSeek API key is not configured.") from None
    except AIServiceError as exc:
        logger.warning("YouTube DOCX: structured content generated no video_id=%s reason=ai_generation_failed", video_id)
        raise YouTubeDocxError("Could not generate DOCX.") from exc
    transcript_short = len(transcript or "") < 3500
    document_options["short_video_note"] = transcript_short

    try:
        file_info = build_docx_file(metadata, transcript, generated_content, target_pages, document_options=document_options)
    except Exception:
        logger.exception("YouTube DOCX: DOCX generated failed video_id=%s", video_id)
        raise YouTubeDocxError("Could not generate DOCX.") from None
    path_exists = bool(file_info.get("path") and file_info["path"].exists())
    logger.info("YouTube DOCX: DOCX created %s video_id=%s", "yes" if path_exists else "no", video_id)
    logger.info("YouTube DOCX: temp path exists %s video_id=%s", "yes" if path_exists else "no", video_id)
    return {
        "download_url": f"/api/youtube-docx/download/{file_info['temp_file_id']}/",
        "title": metadata.get("title") or "YouTube Lecture",
        "estimated_pages": file_info["estimated_pages"],
        "sections_count": sections_count,
        "short_video_note": transcript_short,
        "filename": file_info["filename"],
    }
